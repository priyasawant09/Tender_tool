import fitz
import docx
import os
import re

# parses PDF files
def parse_pdf(file_path: str) -> str:
    try:
        with fitz.open(file_path) as doc:
            # Concatenate text from all pages, clean up whitespace
            text = "\n".join(page.get_text("text").strip() for page in doc)
        return text
    except Exception as e:
        raise ValueError(f"Could not read PDF file '{os.path.basename(file_path)}': {e}")

# parses DOCX files
def parse_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extracting text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extracting text from tables in a readable format
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Could not read DOCX file '{os.path.basename(file_path)}': {e}")

# main parsing function
def parse_cv(file_path: str) -> str:

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file was not found at the specified path: {file_path}")

    _, extension = os.path.splitext(file_path)
    extension = extension.lower()

    if extension == '.pdf':
        return parse_pdf(file_path)
    elif extension == '.docx':
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: '{extension}'. Only PDF and DOCX files are supported.")